from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import TypedDict

import fitz  # pymupdf
from docx import Document as DocxDocument

from helpdesk_rag.exceptions import DocumentLoadError

logger = logging.getLogger(__name__)


class DocumentMetadata(TypedDict, total=False):
    source: str
    format: str
    page_count: int
    headings: list[str]


@dataclass
class Document:
    content: str
    metadata: DocumentMetadata = field(default_factory=dict)


def _extract_markdown_headings(text: str) -> list[str]:
    headings: list[str] = []
    for line in text.splitlines():
        if line.startswith("#"):
            headings.append(line.lstrip("#").strip())
    return headings


def _load_pdf(path: Path) -> Document:
    try:
        doc = fitz.open(path)
    except Exception as exc:
        raise DocumentLoadError(f"Failed to open PDF {path.name}: {exc}") from exc
    try:
        pages = [page.get_text() for page in doc]
    finally:
        doc.close()
    logger.debug("Loaded PDF %s (%d pages)", path.name, len(pages))
    return Document(
        content="\n\n".join(pages),
        metadata={"source": path.name, "format": "pdf", "page_count": len(pages)},
    )


def _load_docx(path: Path) -> Document:
    try:
        doc = DocxDocument(path)
    except Exception as exc:
        raise DocumentLoadError(f"Failed to open DOCX {path.name}: {exc}") from exc
    paragraphs: list[str] = []
    headings: list[str] = []
    for para in doc.paragraphs:
        style = para.style.name.lower() if para.style and para.style.name else ""
        text = para.text.strip()
        if not text:
            continue
        if "heading" in style:
            heading_level = 1
            match = re.search(r"heading\s*(\d)", style)
            if match:
                heading_level = int(match.group(1))
            headings.append(text)
            paragraphs.append(f"{'#' * heading_level} {text}")
        else:
            paragraphs.append(text)
    logger.debug("Loaded DOCX %s (%d paragraphs, %d headings)", path.name, len(paragraphs), len(headings))
    return Document(
        content="\n\n".join(paragraphs),
        metadata={"source": path.name, "format": "docx", "headings": headings},
    )


def _load_markdown(path: Path) -> Document:
    try:
        content = path.read_text(encoding="utf-8")
    except UnicodeDecodeError as exc:
        raise DocumentLoadError(f"Failed to read markdown {path.name}: {exc}") from exc
    headings = _extract_markdown_headings(content)
    logger.debug("Loaded markdown %s (%d headings)", path.name, len(headings))
    return Document(
        content=content.strip(),
        metadata={"source": path.name, "format": "markdown", "headings": headings},
    )


def _load_text(path: Path) -> Document:
    try:
        content = path.read_text(encoding="utf-8")
    except UnicodeDecodeError as exc:
        raise DocumentLoadError(f"Failed to read text {path.name}: {exc}") from exc
    logger.debug("Loaded text %s (%d chars)", path.name, len(content))
    return Document(
        content=content.strip(),
        metadata={"source": path.name, "format": "text"},
    )


_LOADERS = {
    ".pdf": _load_pdf,
    ".docx": _load_docx,
    ".md": _load_markdown,
    ".markdown": _load_markdown,
    ".txt": _load_text,
}

SUPPORTED_EXTENSIONS = set(_LOADERS.keys())


def load_document(path: str | Path) -> Document:
    """Load and parse a document by file extension."""
    path = Path(path)
    if not path.exists():
        raise DocumentLoadError(f"File not found: {path}")
    if not path.is_file():
        raise DocumentLoadError(f"Not a file: {path}")
    ext = path.suffix.lower()
    loader = _LOADERS.get(ext)
    if loader is None:
        raise DocumentLoadError(f"Unsupported file type: {ext}. Supported: {sorted(SUPPORTED_EXTENSIONS)}")
    try:
        return loader(path)
    except DocumentLoadError:
        raise
    except Exception as exc:
        raise DocumentLoadError(f"Failed to load {path.name}: {exc}") from exc